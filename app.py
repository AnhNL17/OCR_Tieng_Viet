import io
import os
import streamlit as st
from PIL import Image

# Kiểm tra & import pdf2image an toàn
try:
    import pdf2image
    HAS_PDF2IMAGE = True
except ImportError:
    HAS_PDF2IMAGE = False

# Kiểm tra & import python-docx an toàn
try:
    import docx
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

from modules import xu_ly_core as backend

# ================== CẤU HÌNH POPPLER ==================
# Nếu chạy trên Streamlit Cloud (Linux), POPPLER_PATH để None để tự nhận diện từ system
import platform

if platform.system() == "Windows":
    POPPLER_PATH = r"C:\Users\HP\Downloads\Release-24.08.0-0\poppler-24.08.0\Library\bin"
else:
    POPPLER_PATH = None  # On Linux / Streamlit Cloud
POPPLER_PATH = r"C:\Users\HP\Downloads\Release-24.08.0-0\poppler-24.08.0\Library\bin"

# Tự động tìm poppler trong thư mục dự án
LOCAL_POPPLER_DIR = os.path.join(os.path.dirname(__file__), "poppler", "bin")
if os.path.exists(LOCAL_POPPLER_DIR):
    POPPLER_PATH = LOCAL_POPPLER_DIR

# ================== PAGE CONFIG ==================
st.set_page_config(
    page_title="OCR Tiếng Việt - Ảnh & PDF",
    layout="wide"
)

# ================== CSS ==================
st.markdown("""
<style>
body {
    background-color: #f9fafb;
}
.main-title {
    font-size: 44px;
    font-weight: 800;
    color: #111827;
}
.sub-title {
    font-size: 18px;
    color: #6b7280;
    margin-bottom: 30px;
}
.section-title {
    font-size: 22px;
    font-weight: 700;
    margin-bottom: 10px;
}
.divider {
    height: 1px;
    background-color: #e5e7eb;
    margin: 20px 0;
}
.stButton>button {
    width: 100%;
    height: 50px;
    border-radius: 12px;
    font-size: 17px;
    font-weight: 700;
    background: linear-gradient(90deg, #2563eb, #3b82f6);
    color: white;
}
.stButton>button:hover {
    background: linear-gradient(90deg, #1d4ed8, #2563eb);
}
textarea {
    font-size: 15px !important;
    line-height: 1.6;
}
</style>
""", unsafe_allow_html=True)

# ================== HEADER ==================
st.markdown('<div class="main-title">OCR Tiếng Việt</div>', unsafe_allow_html=True)
st.markdown(
    '<div class="sub-title">Nhận diện văn bản tiếng Việt từ Ảnh hoặc PDF – tối ưu cho cột & dấu</div>',
    unsafe_allow_html=True
)

# ================== TOP CONTROLS ==================
col_upload, col_psm, col_btn = st.columns([2.5, 1.5, 1])

with col_upload:
    uploaded_file = st.file_uploader(
        "Upload ảnh hoặc PDF (PNG / JPG / JPEG / PDF)",
        type=["png", "jpg", "jpeg", "pdf"],
        label_visibility="collapsed"
    )

with col_psm:
    psm_mode = st.radio(
        "Chế độ PSM",
        (3, 4, 6),
        format_func=lambda x: f"PSM {x}",
        horizontal=True
    )

with col_btn:
    st.markdown("<br>", unsafe_allow_html=True)
    run_ocr = st.button("🚀 Chạy OCR")

st.caption("""
• **PSM 3**: Tự động – tài liệu nhiều cột  
• **PSM 4**: Một cột văn bản  
• **PSM 6**: Nguyên khối (dễ lỗi cột)
""")

st.markdown('<div class="divider"></div>', unsafe_allow_html=True)

# Khởi tạo state lưu kết quả OCR
if "ocr_result" not in st.session_state:
    st.session_state.ocr_result = ""

# ================== MAIN VIEW ==================
left, right = st.columns([1, 1.4], gap="large")

# ----- LEFT: PREVIEW -----
with left:
    st.markdown('<div class="section-title">🖼️ File xem trước</div>', unsafe_allow_html=True)

    images_to_process = []

    if uploaded_file:
        file_ext = uploaded_file.name.split(".")[-1].lower()

        if file_ext == "pdf":
            if not HAS_PDF2IMAGE:
                st.error("⚠️ Chưa cài đặt thư viện `pdf2image`. Vui lòng chạy lệnh: `pip install pdf2image`")
            else:
                try:
                    pdf_bytes = uploaded_file.read()
                    
                    # Chuyển đổi PDF sang danh sách ảnh PIL
                    if POPPLER_PATH and os.path.exists(POPPLER_PATH):
                        images_to_process = pdf2image.convert_from_bytes(pdf_bytes, poppler_path=POPPLER_PATH)
                    else:
                        images_to_process = pdf2image.convert_from_bytes(pdf_bytes)

                    st.success(f"📄 File PDF có **{len(images_to_process)}** trang")

                    # Chọn trang để xem trước
                    page_idx = st.number_input(
                        "Trang xem trước:", 
                        min_value=1, 
                        max_value=len(images_to_process), 
                        value=1
                    ) - 1
                    
                    st.image(images_to_process[page_idx], caption=f"Trang {page_idx + 1}/{len(images_to_process)}", use_container_width=True)

                except pdf2image.exceptions.PDFInfoNotInstalledError:
                    st.error("⚠️ Chưa tìm thấy Poppler tại đường dẫn cấu hình.")
                    st.info(f"Đường dẫn hiện tại: `{POPPLER_PATH}`.")
                except Exception as e:
                    st.error(f"Lỗi khi xử lý file PDF: {e}")
        else:
            images_to_process = [uploaded_file]
            st.image(uploaded_file, use_container_width=True)
    else:
        st.info("⬆️ Upload file ảnh hoặc PDF để hiển thị")

# ----- RIGHT: OCR RESULT & DOWNLOAD -----
with right:
    st.markdown('<div class="section-title">📄 Kết quả OCR</div>', unsafe_allow_html=True)

    if uploaded_file and run_ocr:
        if images_to_process:
            results = []
            with st.spinner(f"Đang xử lý OCR ({len(images_to_process)} trang) với PSM {psm_mode}..."):
                for idx, img in enumerate(images_to_process):
                    try:
                        # CHUYỂN ĐỔI: Nếu là ảnh PIL từ PDF, biến nó thành BytesIO để tương thích với backend
                        if isinstance(img, Image.Image):
                            img_byte_arr = io.BytesIO()
                            img.save(img_byte_arr, format='PNG')
                            img_byte_arr.seek(0)
                            input_img = img_byte_arr
                        else:
                            input_img = img

                        text_page = backend.lay_text_tu_anh(input_img, che_do_doc=psm_mode)
                        
                        if len(images_to_process) > 1:
                            results.append(f"--- TRANG {idx + 1} ---\n{text_page}")
                        else:
                            results.append(text_page)
                    except Exception as e:
                        results.append(f"--- TRANG {idx + 1} ---\nLỗi xử lý: {e}")

                st.session_state.ocr_result = "\n\n".join(results)
        else:
            st.warning("Chưa có dữ liệu hình ảnh để chạy OCR.")

    # Hiển thị ô văn bản kết quả
    output_text = st.text_area(
        "Văn bản nhận diện được",
        value=st.session_state.ocr_result,
        height=500
    )

    # Nút Tải về
    if st.session_state.ocr_result:
        st.markdown("### 💾 Lưu kết quả về máy")
        col_down1, col_down2 = st.columns(2)

        # 1. Tải về file .TXT
        with col_down1:
            st.download_button(
                label="📥 Tải file .TXT",
                data=st.session_state.ocr_result.encode("utf-8"),
                file_name="ket_qua_ocr.txt",
                mime="text/plain; charset=utf-8",
                use_container_width=True
            )

        # 2. Tải về file .DOCX (Word)
        with col_down2:
            if HAS_DOCX:
                doc = docx.Document()
                doc.add_heading("KẾT QUẢ OCR VIETNAMESE", level=1)
                for line in st.session_state.ocr_result.split("\n"):
                    doc.add_paragraph(line)
                
                bio = io.BytesIO()
                doc.save(bio)
                
                st.download_button(
                    label="📥 Tải file .DOCX (Word)",
                    data=bio.getvalue(),
                    file_name="ket_qua_ocr.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True
                )
            else:
                st.warning("Hãy cài `python-docx` để bật tính năng tải file Word: `pip install python-docx`")
    else:
        if not run_ocr:
            st.info("⬅️ Upload file và bấm **Chạy OCR** để xem kết quả")
